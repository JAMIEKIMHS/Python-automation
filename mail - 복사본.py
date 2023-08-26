import zipfile
import glob
import os
from shutil import copyfile, move
from matplotlib import pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.utils import formatdate
from email.mime.application import MIMEApplication
from email.mime.base import MIMEBase
from email.encoders import encode_base64
from email.header import Header



with zipfile.ZipFile("./이름_생년_성별_10000.zip") as f:
    for file in f.infolist():
        file.filename = file.filename.encode('cp437').decode("euc-kr")
        f.extract(file, "./aaaa")

filelist = glob.glob("./aaaa/*.txt")

male = []
female =[] 
if os.path.isdir("./male") == False:
    os.mkdir("./male")
if os.path.isdir("./female") == False:
    os.mkdir("./female")
    
for file in filelist:
    gender = os.path.split(file)[-1].split("_")[-1][0]
    with open(f"{file}", "r") as f:
        if gender in ['1', '3']:
            male.append(f.read())
        else:
            female.append(f.read())

    if gender in ['1', '3']:
        move(file, "./male/")
    else:
        move(file, "./female/")


male_data = sum(map(lambda x : int(x.replace(",", "")), male))
female_data =sum(map(lambda x : int(x.replace(",", "")),female))

plt.bar(['male', 'female'], [male_data, female_data])
plt.title("Agg Money")
plt.savefig("result.png", dpi=100)



document = Document()
document.add_heading("성별에 따른 총 금액 집계")

text = """
- 남녀 집계 정보 
< 그래프>
"""

para = document.add_paragraph()

run = para.add_run(text)

run.bold = True

table = document.add_table(1, 2)

# 테이블 제목 설정 
heading_cells = table.rows[0].cells
heading_cells[0].text = '성별'
heading_cells[1].text = "총금액"

cells = table.add_row().cells

cells[0].text = "남"
cells[1].text = str(male_data)

cells = table.add_row().cells

cells[0].text = "여"
cells[1].text = str(female_data)

table.style = 'LightShading-Accent1'

document.add_picture("./result.png", width=Inches(5.0))

document.save("./보고서.docx")



# id_, pass_ = ("구글 이메일", "비번")

smtp = smtplib.SMTP_SSL("smtp.gmail.com", 465) 
smtp.login(id_, pass_)


email_message_html = """과장님 안녕하세요.. <br>

<h1> 요청하신 내용 첨부합니다. . </h1>

<p><span style="font-weight:bold">감사합니다.</span> </p>

"""


email_from = "보내는 이메일"
email_to = "받는 이메일"
email_date = formatdate(localtime=True)
email_subject = "보고서 자료"


msg = MIMEMultipart('mixed')
msg['From'] = email_from
msg['To'] = email_to
msg['Date'] = email_date
msg['Subject'] = email_subject

# 실제 파일 경로를 open함수로 read 한다.
filename = "보고서.docx"
target_file = "./" + filename
with open(target_file, "rb") as attach_file:
    # application, octet-stream은 모든 종류의 파일을 전송할 때 사용하는 타입 
    file_data = MIMEBase("application", "octet-stream")
    file_data.set_payload(attach_file.read())
    encode_base64(file_data)
    # add_header에 filename를 전달해야 첨부된 파일명이 보입니다. 
    file_data.add_header('Content-Disposition', 'attachment', filename=filename)
    msg.attach(file_data)


msg.attach(MIMEText(email_message_html, "html", _charset="utf-8"))
smtp.sendmail(id_, email_to, msg.as_string())
